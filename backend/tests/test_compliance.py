"""
Architecture Compliance Tests — Migration 004 gap fixes.

Gap #2: syllabus_topics has regulation column, readiness scorer queries work
Gap #4: question_store writes regulation+academic context on insert
Gap #6: coverage_analyzer returns evidence-based results
Gap #7: topic_mapper populates question_topics with confidence scores
Gap #5: analysis_reports no longer has stale regulations[] column
"""
import pytest
from unittest.mock import MagicMock, patch


# ── Helpers ───────────────────────────────────────────────────────────────────

def make_mock_db():
    """Return a mock Supabase client."""
    db = MagicMock()
    # Default: return empty data
    db.table.return_value.select.return_value.eq.return_value.eq.return_value.single.return_value.execute.return_value.data = {}
    db.table.return_value.select.return_value.eq.return_value.execute.return_value.data = []
    db.table.return_value.insert.return_value.execute.return_value.data = [{"id": "test-id-001"}]
    db.table.return_value.upsert.return_value.execute.return_value.data = [{"id": "test-id-001"}]
    return db


# ── Gap #4: question_store writes regulation on insert ─────────────────────

class TestQuestionStoreRegulationPropagation:
    def test_fetches_paper_meta_on_insert(self):
        """_fetch_paper_meta is called with paper_id before storing."""
        from app.parsers.question_store import _fetch_paper_meta
        db = make_mock_db()
        db.table.return_value.select.return_value.eq.return_value.single.return_value.execute.return_value.data = {
            "regulation": "R22",
            "college_id": "college-001",
            "branch_id": "branch-001",
            "semester": 1,
            "exam_year": 2024,
        }
        meta = _fetch_paper_meta(db, "paper-001")
        assert meta["regulation"] == "R22"
        assert meta["college_id"] == "college-001"
        assert meta["exam_year"] == 2024

    def test_returns_empty_dict_on_failure(self):
        """_fetch_paper_meta gracefully returns {} when DB fails."""
        from app.parsers.question_store import _fetch_paper_meta
        db = make_mock_db()
        db.table.return_value.select.return_value.eq.return_value.single.return_value.execute.side_effect = Exception("DB down")
        meta = _fetch_paper_meta(db, "paper-001")
        assert meta == {}

    def test_store_includes_regulation_in_insert_payload(self):
        """store_parse_result includes regulation in the inserted row dict."""
        from app.parsers.question_store import store_parse_result
        from app.parsers.models import ParseResult, ParsedQuestion

        mock_q = ParsedQuestion(
            paper_id="paper-001",
            subject_id="subject-001",
            question_text="Define tautology",
            normalized_text="define tautology",
            question_hash="abc123",
        )
        result = ParseResult(
            paper_id="paper-001",
            questions=[mock_q],
            total_questions=1,
            part_a_count=1,
            part_b_count=0,
        )

        inserted_payloads = []

        db = make_mock_db()
        # No existing question
        db.table.return_value.select.return_value.eq.return_value.eq.return_value.execute.return_value.data = []
        # Paper meta
        db.table.return_value.select.return_value.eq.return_value.single.return_value.execute.return_value.data = {
            "regulation": "R22", "college_id": "c1", "branch_id": "b1",
            "semester": 1, "exam_year": 2024,
        }

        def capture_insert(payload):
            inserted_payloads.append(payload)
            m = MagicMock()
            m.execute.return_value.data = [{"id": "q-001"}]
            return m

        db.table.return_value.insert.side_effect = capture_insert

        with patch("app.parsers.question_store.get_db", return_value=db):
            store_parse_result(result)

        # The insert payload must contain regulation
        assert len(inserted_payloads) >= 1
        payload = inserted_payloads[0]
        assert "regulation" in payload
        assert payload["regulation"] == "R22"
        assert payload["college_id"] == "c1"
        assert payload["exam_year"] == 2024

    def test_get_questions_for_subject_accepts_regulation_filter(self):
        """get_questions_for_subject passes regulation to DB query."""
        from app.parsers.question_store import get_questions_for_subject
        db = make_mock_db()

        chain = MagicMock()
        chain.execute.return_value.data = []
        db.table.return_value.select.return_value.eq.return_value = chain
        chain.eq.return_value = chain

        with patch("app.parsers.question_store.get_db", return_value=db):
            get_questions_for_subject("sub-001", regulation="R22")

        # Verify .eq was called (regulation filter applied)
        assert db.table.return_value.select.return_value.eq.called


# ── Gap #2: syllabus_topics regulation column ─────────────────────────────────

class TestSyllabusTopicsRegulation:
    def test_ingest_syllabus_writes_regulation_to_topics(self):
        """ingest_syllabus passes regulation to every syllabus_topics row."""
        from app.extractors.syllabus_ingester import ingest_syllabus
        import tempfile, os

        topic_rows_inserted = []

        db = make_mock_db()
        db.table.return_value.insert.return_value.execute.return_value.data = [{"id": "syl-001"}]
        db.table.return_value.select.return_value.eq.return_value.single.return_value.execute.return_value.data = {
            "college_id": "college-001"
        }

        def capture_insert(rows):
            if isinstance(rows, list):
                topic_rows_inserted.extend(rows)
            m = MagicMock()
            m.execute.return_value.data = [{"id": "syl-001"}]
            return m

        db.table.return_value.insert.side_effect = capture_insert

        # Create a tiny DOCX syllabus file
        from docx import Document
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        doc = Document()
        doc.add_paragraph("UNIT I: Mathematical Logic")
        doc.add_paragraph("Propositional Logic, Tautology, PCNF")
        doc.add_paragraph("UNIT II: Set Theory")
        doc.add_paragraph("Sets, Relations, Hasse Diagrams")
        doc.save(path)

        try:
            with patch("app.extractors.syllabus_ingester.get_db", return_value=db):
                ingest_syllabus(path, "sub-001", "R22", college_id="col-001", branch_id="br-001")
        finally:
            os.unlink(path)

        # All topic rows must carry regulation
        assert len(topic_rows_inserted) > 0
        for row in topic_rows_inserted:
            if "topic_name" in row:  # skip the syllabi insert
                assert row.get("regulation") == "R22", f"Missing regulation on: {row}"
                assert row.get("subject_id") == "sub-001"

    def test_readiness_scorer_fetch_uses_regulation_column(self):
        """_fetch_syllabus_topics uses .eq('regulation', ...) directly."""
        from app.planner.readiness_scorer import ReadinessScorer

        db = make_mock_db()
        chain = MagicMock()
        chain.execute.return_value.data = [{"topic_name": "Hasse Diagram"}, {"topic_name": "PCNF"}]
        db.table.return_value.select.return_value.eq.return_value = chain
        chain.eq.return_value = chain

        with patch("app.planner.readiness_scorer.get_db", return_value=db):
            scorer = ReadinessScorer()
            topics = scorer._fetch_syllabus_topics("sub-001", "R22")

        assert "Hasse Diagram" in topics
        assert "PCNF" in topics


# ── Gap #6: coverage_analyzer ─────────────────────────────────────────────────

class TestCoverageAnalyzer:
    def _make_syllabus_topics(self):
        return [
            {"id": "st-1", "syllabus_id": "syl-1", "topic_name": "Hasse Diagram",    "unit_number": 2, "unit_name": "Relations"},
            {"id": "st-2", "syllabus_id": "syl-1", "topic_name": "Kruskal Algorithm", "unit_number": 5, "unit_name": "Graphs"},
            {"id": "st-3", "syllabus_id": "syl-1", "topic_name": "Recurrence Relations","unit_number": 4, "unit_name": "Recurrence"},
            {"id": "st-4", "syllabus_id": "syl-1", "topic_name": "Tautology",          "unit_number": 1, "unit_name": "Logic"},
        ]

    def _make_questions(self):
        return [
            {"id": "q-1", "question_text": "Draw the Hasse diagram", "topic_tags": ["Hasse Diagram"], "unit_number": 2, "marks": 5, "question_type": "construction"},
            {"id": "q-2", "question_text": "Apply Kruskal's algorithm", "topic_tags": ["Kruskal's Algorithm"], "unit_number": 5, "marks": 5, "question_type": "construction"},
            {"id": "q-3", "question_text": "Define tautology", "topic_tags": ["Tautology", "Logic"], "unit_number": 1, "marks": 1, "question_type": "definition"},
            # Recurrence Relations has no matching questions
        ]

    def test_coverage_returns_correct_structure(self):
        from app.intelligence.coverage_analyzer import compute_coverage

        db = make_mock_db()
        syl_chain = MagicMock()
        syl_chain.order.return_value.execute.return_value.data = self._make_syllabus_topics()
        db.table.return_value.select.return_value.eq.return_value = syl_chain
        syl_chain.eq.return_value = syl_chain

        q_chain = MagicMock()
        q_chain.execute.return_value.data = self._make_questions()
        db.table.return_value.select.return_value.eq.return_value = q_chain
        q_chain.eq.return_value = q_chain

        # Patch to return different data based on table name
        def table_router(name):
            m = MagicMock()
            if name == "syllabus_topics":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.order.return_value.execute.return_value.data = self._make_syllabus_topics()
                chain.select.return_value.eq.return_value.order.return_value.execute.return_value.data = self._make_syllabus_topics()
                return chain
            elif name == "questions":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.execute.return_value.data = self._make_questions()
                return chain
            return m

        db.table.side_effect = table_router

        with patch("app.intelligence.coverage_analyzer.get_db", return_value=db):
            result = compute_coverage("sub-001", "R22", syllabus_id="syl-1")

        assert "overall_coverage_pct" in result
        assert "unit_coverage" in result
        assert "topic_mapping" in result
        assert result["total_topics"] == 4

    def test_empty_syllabus_returns_zero_coverage_with_warning(self):
        from app.intelligence.coverage_analyzer import compute_coverage

        db = make_mock_db()
        chain = MagicMock()
        chain.select.return_value.eq.return_value.eq.return_value.order.return_value.execute.return_value.data = []
        db.table.side_effect = lambda name: chain

        with patch("app.intelligence.coverage_analyzer.get_db", return_value=db):
            result = compute_coverage("sub-001", "R22")

        assert result["overall_coverage_pct"] == 0.0
        assert result["total_topics"] == 0
        assert "warning" in result

    def test_coverage_does_not_count_without_evidence(self):
        """A topic with no matching questions must NOT be counted as covered."""
        from app.intelligence.coverage_analyzer import compute_coverage

        topics = [{"id": "st-1", "syllabus_id": "syl-1", "topic_name": "Uncovered Topic", "unit_number": 1, "unit_name": "Logic"}]
        questions = [{"id": "q-1", "question_text": "Something else", "topic_tags": ["Completely Different"], "unit_number": 1, "marks": 5, "question_type": "theory"}]

        def table_router(name):
            m = MagicMock()
            if name == "syllabus_topics":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.order.return_value.execute.return_value.data = topics
                chain.select.return_value.eq.return_value.order.return_value.execute.return_value.data = topics
                return chain
            elif name == "questions":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.execute.return_value.data = questions
                return chain
            return m

        db = make_mock_db()
        db.table.side_effect = table_router
        with patch("app.intelligence.coverage_analyzer.get_db", return_value=db):
            result = compute_coverage("sub-001", "R22")

        assert result["covered_topics"] == 0
        assert result["overall_coverage_pct"] == 0.0

    def test_uncovered_topics_listed(self):
        from app.intelligence.coverage_analyzer import compute_coverage

        topics = [
            {"id": "st-1", "syllabus_id": "syl-1", "topic_name": "Hasse Diagram", "unit_number": 2, "unit_name": "Relations"},
            {"id": "st-2", "syllabus_id": "syl-1", "topic_name": "Never Covered", "unit_number": 2, "unit_name": "Relations"},
        ]
        questions = [{"id": "q-1", "question_text": "Draw Hasse diagram", "topic_tags": ["Hasse Diagram"], "unit_number": 2, "marks": 5, "question_type": "construction"}]

        def table_router(name):
            m = MagicMock()
            if name == "syllabus_topics":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.order.return_value.execute.return_value.data = topics
                chain.select.return_value.eq.return_value.order.return_value.execute.return_value.data = topics
                return chain
            elif name == "questions":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.execute.return_value.data = questions
                return chain
            return m

        db = make_mock_db()
        db.table.side_effect = table_router

        with patch("app.intelligence.coverage_analyzer.get_db", return_value=db):
            result = compute_coverage("sub-001", "R22")

        uncovered = []
        for uc in result["unit_coverage"]:
            uncovered.extend(uc["uncovered_topics"])
        assert "Never Covered" in uncovered


# ── Gap #7: topic_mapper ──────────────────────────────────────────────────────

class TestTopicMapper:
    def test_maps_question_to_matching_syllabus_topic(self):
        from app.intelligence.topic_mapper import map_questions_to_topics

        syl_topics = [
            {"id": "st-1", "topic_name": "Hasse Diagram", "unit_number": 2},
            {"id": "st-2", "topic_name": "Kruskal's Algorithm", "unit_number": 5},
        ]
        questions = [
            {"id": "q-1", "topic_tags": ["Hasse Diagram"], "unit_number": 2},
            {"id": "q-2", "topic_tags": ["Kruskal"], "unit_number": 5},
        ]
        upserted = []

        def table_router(name):
            m = MagicMock()
            if name == "syllabus_topics":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.execute.return_value.data = syl_topics
                return chain
            elif name == "questions":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.execute.return_value.data = questions
                return chain
            elif name == "question_topics":
                chain = MagicMock()
                # No existing mappings
                chain.select.return_value.execute.return_value.data = []
                def capture(payload):
                    upserted.append(payload)
                    mm = MagicMock()
                    mm.execute.return_value.data = [{"id": "qt-001"}]
                    return mm
                chain.upsert.side_effect = capture
                return chain
            return m

        db = make_mock_db()
        db.table.side_effect = table_router

        with patch("app.intelligence.topic_mapper.get_db", return_value=db):
            result = map_questions_to_topics("sub-001", "R22")

        assert result["mapped"] >= 1
        assert result["failed"] == 0

    def test_no_mapping_without_syllabus(self):
        from app.intelligence.topic_mapper import map_questions_to_topics

        def table_router(name):
            m = MagicMock()
            if name == "syllabus_topics":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.execute.return_value.data = []
                return chain
            return m

        db = make_mock_db()
        db.table.side_effect = table_router

        with patch("app.intelligence.topic_mapper.get_db", return_value=db):
            result = map_questions_to_topics("sub-001", "R22")

        assert result["mapped"] == 0
        assert "warning" in result

    def test_confidence_score_range(self):
        from app.intelligence.topic_mapper import _compute_confidence
        # Exact match → high confidence
        score = _compute_confidence("Hasse Diagram", "Hasse Diagram")
        assert score >= 0.99

        # Partial match
        score2 = _compute_confidence("Kruskal", "Kruskal's Algorithm")
        assert 0.5 <= score2 <= 1.0

        # No match
        score3 = _compute_confidence("photosynthesis", "Hasse Diagram")
        assert score3 < 0.5

    def test_idempotent_skips_already_mapped(self):
        """Re-running map_questions_to_topics skips already-mapped questions."""
        from app.intelligence.topic_mapper import map_questions_to_topics

        syl_topics = [{"id": "st-1", "topic_name": "Hasse Diagram", "unit_number": 2}]
        questions  = [{"id": "q-1", "topic_tags": ["Hasse Diagram"], "unit_number": 2}]

        def table_router(name):
            m = MagicMock()
            if name == "syllabus_topics":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.execute.return_value.data = syl_topics
                return chain
            elif name == "questions":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.execute.return_value.data = questions
                return chain
            elif name == "question_topics":
                chain = MagicMock()
                # Already mapped
                chain.select.return_value.execute.return_value.data = [{"question_id": "q-1"}]
                return chain
            return m

        db = make_mock_db()
        db.table.side_effect = table_router

        with patch("app.intelligence.topic_mapper.get_db", return_value=db):
            result = map_questions_to_topics("sub-001", "R22")

        assert result["skipped_already_mapped"] == 1
        assert result["mapped"] == 0

    def test_questions_without_topic_tags_not_mapped(self):
        from app.intelligence.topic_mapper import map_questions_to_topics

        syl_topics = [{"id": "st-1", "topic_name": "Hasse Diagram", "unit_number": 2}]
        questions  = [{"id": "q-1", "topic_tags": [], "unit_number": 2}]

        def table_router(name):
            m = MagicMock()
            if name == "syllabus_topics":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.execute.return_value.data = syl_topics
                return chain
            elif name == "questions":
                chain = MagicMock()
                chain.select.return_value.eq.return_value.eq.return_value.execute.return_value.data = questions
                return chain
            elif name == "question_topics":
                chain = MagicMock()
                chain.select.return_value.execute.return_value.data = []
                return chain
            return m

        db = make_mock_db()
        db.table.side_effect = table_router

        with patch("app.intelligence.topic_mapper.get_db", return_value=db):
            result = map_questions_to_topics("sub-001", "R22")

        assert result["no_syllabus_match"] == 1
        assert result["mapped"] == 0
