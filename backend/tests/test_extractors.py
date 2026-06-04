"""
Milestone 2 extractor tests.
Uses real temp files — no mocks for the extraction logic itself.
"""
import os
import pytest
import tempfile


# ── Helpers to create test files ─────────────────────────────────────────────

def make_pdf(path: str, text: str = "Hello PDF World\nUnit 1\nUnit 2") -> str:
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 100), text)
    doc.save(path)
    doc.close()
    return path


def make_docx(path: str) -> str:
    from docx import Document
    doc = Document()
    doc.add_paragraph("Course Code: A6CS08")
    doc.add_paragraph("DISCRETE MATHEMATICS")
    # Add a table like MLRIT papers
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "1"
    table.cell(0, 1).text = "a)"
    table.cell(0, 2).text = "Define Tautology and Contradiction"
    table.cell(1, 0).text = "2"
    table.cell(1, 1).text = "a)"
    table.cell(1, 2).text = "Draw the Hasse diagram for positive divisors of 45"
    table.cell(2, 0).text = "3"
    table.cell(2, 1).text = "a)"
    table.cell(2, 2).text = "Apply Kruskal's algorithm on the following graph"
    doc.save(path)
    return path


def make_html(path: str) -> str:
    html = "<html><body><h1>Unit 1</h1><p>Propositional Logic</p></body></html>"
    with open(path, "w") as f:
        f.write(html)
    return path


def make_zip(path: str, files: dict) -> str:
    import zipfile
    with zipfile.ZipFile(path, "w") as z:
        for name, content in files.items():
            z.writestr(name, content)
    return path


# ── PDF Extractor ─────────────────────────────────────────────────────────────

class TestPDFExtractor:
    def test_extracts_text(self, tmp_path):
        from app.extractors.pdf_extractor import PDFExtractor
        pdf = make_pdf(str(tmp_path / "test.pdf"), "Hello PDF\nUnit 1 Logic")
        extractor = PDFExtractor()
        assert extractor.can_handle(pdf)
        doc = extractor.extract(pdf)
        assert "Hello PDF" in doc.raw_text
        assert doc.file_type == "pdf"
        assert doc.page_count == 1
        assert doc.file_size_bytes > 0

    def test_cannot_handle_docx(self, tmp_path):
        from app.extractors.pdf_extractor import PDFExtractor
        assert not PDFExtractor().can_handle("file.docx")

    def test_raises_on_missing_file(self):
        from app.extractors.pdf_extractor import PDFExtractor
        from app.extractors.base import ExtractionError
        with pytest.raises(ExtractionError):
            PDFExtractor().extract("/nonexistent/path/file.pdf")


# ── DOCX Extractor ────────────────────────────────────────────────────────────

class TestDocxExtractor:
    def test_extracts_paragraphs(self, tmp_path):
        from app.extractors.docx_extractor import DocxExtractor
        path = make_docx(str(tmp_path / "paper.docx"))
        extractor = DocxExtractor()
        assert extractor.can_handle(path)
        doc = extractor.extract(path)
        assert "DISCRETE MATHEMATICS" in doc.raw_text
        assert doc.file_type == "docx"

    def test_extracts_table_cells(self, tmp_path):
        from app.extractors.docx_extractor import DocxExtractor
        path = make_docx(str(tmp_path / "paper.docx"))
        doc = DocxExtractor().extract(path)
        # Table questions must be in raw_text
        assert "Tautology" in doc.raw_text
        assert "Hasse diagram" in doc.raw_text
        assert "Kruskal" in doc.raw_text

    def test_metadata_has_table_count(self, tmp_path):
        from app.extractors.docx_extractor import DocxExtractor
        path = make_docx(str(tmp_path / "paper.docx"))
        doc = DocxExtractor().extract(path)
        assert doc.metadata.get("tables", 0) >= 1

    def test_cannot_handle_pdf(self):
        from app.extractors.docx_extractor import DocxExtractor
        assert not DocxExtractor().can_handle("file.pdf")


# ── HTML Extractor ────────────────────────────────────────────────────────────

class TestHtmlExtractor:
    def test_extracts_text(self, tmp_path):
        from app.extractors.html_extractor import HtmlExtractor
        path = make_html(str(tmp_path / "page.html"))
        extractor = HtmlExtractor()
        assert extractor.can_handle(path)
        doc = extractor.extract(path)
        assert "Unit 1" in doc.raw_text
        assert "Propositional Logic" in doc.raw_text
        assert doc.file_type == "html"

    def test_strips_script_tags(self, tmp_path):
        from app.extractors.html_extractor import HtmlExtractor
        path = str(tmp_path / "page.html")
        with open(path, "w") as f:
            f.write("<html><body><script>alert('x')</script><p>Content</p></body></html>")
        doc = HtmlExtractor().extract(path)
        assert "alert" not in doc.raw_text
        assert "Content" in doc.raw_text


# ── Archive Extractor ─────────────────────────────────────────────────────────

class TestArchiveExtractor:
    def test_extracts_zip(self, tmp_path):
        from app.extractors.archive_extractor import ArchiveExtractor
        zip_path = make_zip(
            str(tmp_path / "papers.zip"),
            {"paper1.txt": "Unit 1 content", "paper2.txt": "Unit 2 content"},
        )
        extractor = ArchiveExtractor()
        assert extractor.can_handle(zip_path)
        out_dir = str(tmp_path / "extracted")
        files = extractor.extract_to_dir(zip_path, out_dir)
        assert len(files) == 2
        assert any("paper1.txt" in f for f in files)

    def test_can_handle_rar(self):
        from app.extractors.archive_extractor import ArchiveExtractor
        assert ArchiveExtractor().can_handle("archive.rar")

    def test_raises_on_missing_archive(self):
        from app.extractors.archive_extractor import ArchiveExtractor
        from app.extractors.base import ExtractionError
        with pytest.raises(ExtractionError):
            ArchiveExtractor().extract_to_dir("/nonexistent.rar", "/tmp/out")


# ── Extractor Factory ─────────────────────────────────────────────────────────

class TestExtractorFactory:
    def test_routes_pdf(self, tmp_path):
        from app.extractors.extractor_factory import extract
        path = make_pdf(str(tmp_path / "test.pdf"))
        doc = extract(path)
        assert doc.file_type == "pdf"
        assert isinstance(doc.raw_text, str)

    def test_routes_docx(self, tmp_path):
        from app.extractors.extractor_factory import extract
        path = make_docx(str(tmp_path / "test.docx"))
        doc = extract(path)
        assert doc.file_type == "docx"

    def test_routes_html(self, tmp_path):
        from app.extractors.extractor_factory import extract
        path = make_html(str(tmp_path / "test.html"))
        doc = extract(path)
        assert doc.file_type == "html"

    def test_raises_on_unsupported_type(self, tmp_path):
        from app.extractors.extractor_factory import extract
        from app.extractors.base import ExtractionError
        path = str(tmp_path / "file.xyz")
        with open(path, "w") as f:
            f.write("data")
        with pytest.raises(ExtractionError):
            extract(path)

    def test_archive_and_process_zip(self, tmp_path):
        from app.extractors.extractor_factory import extract_archive_and_process
        # Create a zip with a docx inside
        docx_path = make_docx(str(tmp_path / "paper.docx"))
        zip_path = str(tmp_path / "bundle.zip")
        import zipfile
        with zipfile.ZipFile(zip_path, "w") as z:
            z.write(docx_path, "paper.docx")
        out_dir = str(tmp_path / "out")
        docs = extract_archive_and_process(zip_path, out_dir)
        assert len(docs) == 1
        assert docs[0].file_type == "docx"
        assert "DISCRETE MATHEMATICS" in docs[0].raw_text


# ── Text Normalizer ───────────────────────────────────────────────────────────

class TestTextNormalizer:
    def test_collapses_blank_lines(self):
        from app.extractors.text_normalizer import normalize
        # 5 blank lines → at most 2 blank lines (≤ 3 newlines between content)
        text = "Line 1\n\n\n\n\nLine 2"
        result = normalize(text)
        assert "\n\n\n\n" not in result  # never more than 2 consecutive blanks

    def test_strips_word_artifacts(self):
        from app.extractors.text_normalizer import normalize
        text = "~$tempfile\nReal content"
        result = normalize(text)
        assert "~$" not in result
        assert "Real content" in result

    def test_normalizes_em_dash(self):
        from app.extractors.text_normalizer import normalize
        result = normalize("Unit 1 — Logic")
        assert "-" in result
        assert "—" not in result


# ── Syllabus Parser ───────────────────────────────────────────────────────────

class TestSyllabusParser:
    def test_parses_units(self):
        from app.extractors.syllabus_ingester import parse_syllabus_text
        text = """
UNIT I: Mathematical Logic
Propositional logic, Predicates, Quantifiers

UNIT II: Set Theory and Relations
Sets, Relations, Functions, Hasse Diagrams, Lattices

UNIT III: Combinatorics
Permutations, Combinations, Pigeonhole Principle

UNIT IV: Recurrence Relations
Characteristic roots, Generating functions

UNIT V: Graph Theory
Trees, Spanning trees, Kruskal's algorithm, Chromatic number
"""
        units = parse_syllabus_text(text)
        assert len(units) == 5
        assert units[0].unit_name == "Mathematical Logic"
        assert units[1].unit_number == 2
        assert any("Hasse" in t.topic_name for t in units[1].topics)

    def test_handles_empty_text(self):
        from app.extractors.syllabus_ingester import parse_syllabus_text
        assert parse_syllabus_text("") == []

    def test_handles_roman_numerals(self):
        from app.extractors.syllabus_ingester import parse_syllabus_text
        text = "UNIT III: Combinatorics\nTopics here"
        units = parse_syllabus_text(text)
        assert units[0].unit_number == 3
