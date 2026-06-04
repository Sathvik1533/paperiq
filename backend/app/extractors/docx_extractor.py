import os
from docx import Document
from app.extractors.base import BaseExtractor, ExtractedDocument, ExtractionError
from app.logger import get_logger

log = get_logger(__name__)


class DocxExtractor(BaseExtractor):
    """
    Extracts text from DOCX files using python-docx.
    Reads both paragraph text AND table cell text — critical for
    MLRIT question papers which store questions inside tables.
    """

    def can_handle(self, file_path: str) -> bool:
        return file_path.lower().endswith(".docx")

    def extract(self, file_path: str) -> ExtractedDocument:
        if not os.path.exists(file_path):
            raise ExtractionError(f"File not found: {file_path}")
        try:
            doc = Document(file_path)
            parts = []

            # Paragraphs
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    parts.append(t)

            # Tables — question papers store questions in table rows
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    # Deduplicate merged cells (python-docx repeats merged cells)
                    seen = []
                    for c in row_cells:
                        if c not in seen:
                            seen.append(c)
                    if seen:
                        parts.append(" | ".join(seen))

            raw_text = "\n".join(parts).strip()
            log.info(f"[DocxExtractor] Extracted {len(parts)} text blocks from {os.path.basename(file_path)}")
            return ExtractedDocument(
                raw_text=raw_text,
                file_path=file_path,
                file_type="docx",
                file_name=os.path.basename(file_path),
                file_size_bytes=os.path.getsize(file_path),
                metadata={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
                extraction_method="python-docx",
            )
        except Exception as e:
            raise ExtractionError(f"DOCX extraction failed for {file_path}: {e}") from e
